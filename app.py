import streamlit as st
import pandas as pd
from docx import Document

st.set_page_config(page_title="Universal File → CSV Converter", page_icon="📄")

st.title("📄 XLS / TXT / DOC / DOCX / CSV → CSV Converter (Semicolon Separated)")

uploaded_file = st.file_uploader(
    "Upload XLS, XLSX, TXT, DOC, DOCX, or CSV file",
    type=["xls", "xlsx", "txt", "doc", "docx", "csv"]
)

def read_word_file(file):
    doc = Document(file)

    # Try tables first
    if doc.tables:
        table = doc.tables[0]
        data = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            data.append(cells)
        return pd.DataFrame(data)

    # Fallback: text paragraphs
    text_lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    if not text_lines:
        return pd.DataFrame()

    rows = [line.split() for line in text_lines]
    return pd.DataFrame(rows)

if uploaded_file:
    file_name = uploaded_file.name.lower()

    # ---- EXCEL FILES ----
    if file_name.endswith((".xls", ".xlsx")):
        df = pd.read_excel(uploaded_file)

    # ---- TXT FILES ----
    elif file_name.endswith(".txt"):
        content = uploaded_file.read().decode("utf-8", errors="replace")

        if "\t" in content:
            sep = "\t"
        elif ";" in content:
            sep = ";"
        elif "," in content:
            sep = ","
        else:
            sep = " "

        uploaded_file.seek(0)
        df = pd.read_csv(uploaded_file, sep=sep)

    # ---- CSV FILES ----
    elif file_name.endswith(".csv"):
        content = uploaded_file.read().decode("utf-8", errors="replace")

        if ";" in content:
            sep = ";"
        elif "\t" in content:
            sep = "\t"
        elif "," in content:
            sep = ","
        else:
            sep = ","  # fallback

        uploaded_file.seek(0)
        df = pd.read_csv(uploaded_file, sep=sep)

    # ---- WORD FILES ----
    elif file_name.endswith((".doc", ".docx")):
        df = read_word_file(uploaded_file)

    # ---- CONVERSION TO CSV (semicolon) ----
    csv_data = df.to_csv(index=False, sep=";").encode("utf-8")

    st.success("File converted successfully!")

    st.download_button(
        label="⬇️ Download CSV file",
        data=csv_data,
        file_name=file_name.replace(".xls", ".csv")
                           .replace(".xlsx", ".csv")
                           .replace(".txt", ".csv")
                           .replace(".doc", ".csv")
                           .replace(".docx", ".csv"),
        mime="text/csv",
    )
